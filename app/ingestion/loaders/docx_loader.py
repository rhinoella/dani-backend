"""
DOCX document loader using python-docx.

Extracts text content from Word documents with paragraph-level structure.
"""

from __future__ import annotations

import io
import logging
from typing import Dict, List, Any, Optional
from dataclasses import dataclass

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)


@dataclass
class DocxParagraph:
    """Represents a paragraph from a DOCX document."""
    index: int
    text: str
    style: Optional[str]
    is_heading: bool


class DOCXLoader:
    """
    DOCX document loader using python-docx.
    
    Extracts text from Word documents, preserving paragraph structure
    and heading hierarchy for better context.
    """
    
    def __init__(self) -> None:
        self.supported_extensions = [".docx", ".doc"]
        self.heading_styles = {
            "Heading 1", "Heading 2", "Heading 3", 
            "Heading 4", "Heading 5", "Heading 6",
            "Title", "Subtitle"
        }
    
    def load(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Load and extract text from a DOCX file.
        
        Args:
            file_content: Raw bytes of the DOCX file
            filename: Original filename for metadata
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        logger.info(f"Loading DOCX: {filename}")
        
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            
            # Extract document metadata
            doc_metadata = self._extract_metadata(doc)
            doc_metadata["filename"] = filename
            
            # Extract text from paragraphs
            paragraphs: List[DocxParagraph] = []
            full_text_parts: List[str] = []
            current_section: List[str] = []
            current_heading: Optional[str] = None
            
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                style_name = para.style.name if para.style else None
                is_heading = style_name in self.heading_styles if style_name else False
                
                paragraphs.append(DocxParagraph(
                    index=idx,
                    text=text,
                    style=style_name,
                    is_heading=is_heading
                ))
                
                # Format headings differently
                if is_heading:
                    if current_section:
                        full_text_parts.append("\n".join(current_section))
                        current_section = []
                    current_heading = text
                    full_text_parts.append(f"\n## {text}\n")
                else:
                    current_section.append(text)
            
            # Add remaining section
            if current_section:
                full_text_parts.append("\n".join(current_section))
            
            full_text = "\n".join(full_text_parts)
            
            # Extract tables if present
            tables_text = self._extract_tables(doc)
            if tables_text:
                full_text += f"\n\n## Tables\n{tables_text}"
            
            doc_metadata["paragraph_count"] = len(paragraphs)
            doc_metadata["table_count"] = len(doc.tables)
            
            logger.info(f"Extracted {len(paragraphs)} paragraphs from DOCX")
            
            return {
                "text": full_text.strip(),
                "paragraphs": [
                    {
                        "index": p.index,
                        "text": p.text,
                        "style": p.style,
                        "is_heading": p.is_heading
                    }
                    for p in paragraphs
                ],
                "metadata": doc_metadata,
                "paragraph_count": len(paragraphs),
            }
            
        except PackageNotFoundError as e:
            logger.error(f"Invalid DOCX file {filename}: {e}")
            raise ValueError(f"Invalid or corrupted DOCX file: {str(e)}") from e
        except Exception as e:
            logger.error(f"Failed to load DOCX {filename}: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}") from e
    
    def _extract_metadata(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract metadata from DOCX document."""
        metadata: Dict[str, Any] = {}
        
        try:
            core_props = doc.core_properties
            
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["creation_date"] = str(core_props.created)
            if core_props.modified:
                metadata["modification_date"] = str(core_props.modified)
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.comments:
                metadata["comments"] = core_props.comments
                
        except Exception as e:
            logger.warning(f"Could not extract DOCX metadata: {e}")
        
        return metadata
    
    def _extract_tables(self, doc: DocxDocument) -> str:
        """Extract text from tables in the document."""
        if not doc.tables:
            return ""
        
        tables_text: List[str] = []
        
        for table_idx, table in enumerate(doc.tables, start=1):
            rows_text: List[str] = []
            
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            
            if rows_text:
                table_content = f"Table {table_idx}:\n" + "\n".join(rows_text)
                tables_text.append(table_content)
        
        return "\n\n".join(tables_text)
    
    def load_sections(self, file_content: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Load DOCX and return section-based records for chunking.
        
        Sections are determined by headings - each heading starts a new section.
        
        Args:
            file_content: Raw bytes of the DOCX file
            filename: Original filename for metadata
            
        Returns:
            List of section records suitable for chunking
        """
        result = self.load(file_content, filename)
        paragraphs = result["paragraphs"]
        
        if not paragraphs:
            return []
        
        # Group paragraphs by section (heading-based)
        records = []
        current_section: List[str] = []
        current_heading: Optional[str] = None
        section_idx = 0
        
        for para in paragraphs:
            if para["is_heading"]:
                # Save previous section
                if current_section:
                    section_text = "\n".join(current_section)
                    records.append({
                        "text": section_text,
                        "metadata": {
                            "source_file": filename,
                            "section_index": section_idx,
                            "section_heading": current_heading,
                            **result["metadata"],
                        }
                    })
                    section_idx += 1
                
                current_heading = para["text"]
                current_section = []
            else:
                current_section.append(para["text"])
        
        # Add final section
        if current_section:
            section_text = "\n".join(current_section)
            records.append({
                "text": section_text,
                "metadata": {
                    "source_file": filename,
                    "section_index": section_idx,
                    "section_heading": current_heading,
                    **result["metadata"],
                }
            })
        
        # If no sections were created, create one from full text
        if not records and result["text"]:
            records.append({
                "text": result["text"],
                "metadata": {
                    "source_file": filename,
                    "section_index": 0,
                    **result["metadata"],
                }
            })
        
        return records
